import os
import gradio as gr
import numpy as np
import faiss
import json
import re
import docx
import shutil
import zipfile
import logging
from pathlib import Path
from sentence_transformers import SentenceTransformer
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama import OllamaLLM
from typing import Dict, Any
from sentence_transformers import CrossEncoder
import process_files 
from langchain.text_splitter import RecursiveCharacterTextSplitter # 保留，用於對 O.json 中的內容進行分塊

# --- 日誌設定 ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# --- 全域設定 ---
SOURCE_DIR = Path("./data")
DB_DIR = Path("./db")
TEMPLATE_DIR = Path("./templates")
OUTPUT_DIR = Path("./output")
SOURCE_DIR.mkdir(exist_ok=True); DB_DIR.mkdir(exist_ok=True); TEMPLATE_DIR.mkdir(exist_ok=True); OUTPUT_DIR.mkdir(exist_ok=True)
MODEL_NAME = "BAAI/bge-m3"
#OLLAMA_MODEL = "qwen2.5vl:7b" # 您選擇的 LLM 模型
OLLAMA_MODEL = "gpt-oss:20b" # 您選擇的 LLM 模型
RAG_COMPONENTS: Dict[str, Any] = {"model": None, "index": None, "metadata": None}
reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')

# --- 核心邏輯函式 ---

def query_single_fact(question: str, project_scope: str) -> str:
    logger.info(f"開始查詢單一事實：'{question}'，專案範圍：'{project_scope}'")
    index, metadata, model = RAG_COMPONENTS["index"], RAG_COMPONENTS["metadata"], RAG_COMPONENTS["model"]
    
    if not all(RAG_COMPONENTS.values()):
        logger.warning("知識庫組件尚未完整載入，無法執行查詢。")
        return "[錯誤：知識庫組件尚未完整載入，請重新建立索引。]"

    if hasattr(index, 'nprobe'): index.nprobe = 16
    
    try:
        query_embedding = model.encode([question], convert_to_numpy=True)
        if query_embedding.shape[1] != index.d:
            raise ValueError(f"查詢向量維度 {query_embedding.shape[1]} 與索引維度 {index.d} 不匹配。")
        distances, indices = index.search(query_embedding, 5)
        logger.info(f"成功在 FAISS 索引中檢索到 {len(indices[0])} 個候選文件。")
    except Exception as e:
        logger.error(f"檢索發生錯誤：{e}", exc_info=True)
        return f"[檢索發生錯誤：{e}]"

    filtered_docs = [metadata[j] for j in indices[0] if j != -1 and metadata[j].get("project") == project_scope]
    if not filtered_docs:
        logger.warning(f"在專案 '{project_scope}' 中找不到與問題相關的資訊。")
        return f"[在專案 {project_scope} 中找不到相關資訊]"
    
    context = "\n\n".join([doc["content"] for doc in filtered_docs])
    
    prompt_template = (
        "你是一個資料提取 API。你的唯一任務是根據上下文，為問題以繁體中文提供最精確、最原始的答案。不要添加任何標籤、前綴、解釋或思考過程的文字以及任何額外的文字。\n\n"
        "請根據以下上下文回答問題。\n\n"
        "上下文：{context}\n\n"
        "問題：{question}\n"
        "精確答案："
    )
    prompt = ChatPromptTemplate.from_template(prompt_template)
    
    llm = OllamaLLM(model=OLLAMA_MODEL)
    chain = prompt | llm
    answer = chain.invoke({"context": context, "question": question})
    logger.info(f"LLM 成功生成答案。")
    return answer.strip()

def build_knowledge_base(progress=gr.Progress(track_tqdm=True)):
    global RAG_COMPONENTS
    logger.info("開始建立/更新知識庫。")
    
    if RAG_COMPONENTS["model"] is None:
        try:
            progress(0, desc="步驟 1/5: 載入嵌入模型...")
            RAG_COMPONENTS["model"] = SentenceTransformer(MODEL_NAME, trust_remote_code=True)
            logger.info("SentenceTransformer 模型載入成功。")
        except Exception as e:
            logger.error(f"❌ 載入嵌入模型失敗：{e}", exc_info=True)
            return f"❌ 載入模型失敗：{e}"

    # --- 新增步驟：呼叫 process_files.py 來處理原始文件並生成 O.json ---
    progress(0.1, desc="步驟 1.5/5: 呼叫 process_files.py 處理原始文件...")
    try:
        process_files.main() # 執行 process_files.py 中的主函數，它會將結果儲存到 db/O.json
        logger.info("process_files.py 成功執行，db/O.json 已生成。")
    except Exception as e:
        logger.error(f"❌ 呼叫 process_files.py 失敗：{e}", exc_info=True)
        return f"❌ 原始文件處理失敗：{e}"
    # --- 結束新增步驟 ---

    all_chunks_with_metadata = []
    progress(0.2, desc="步驟 2/5: 從 db/O.json 讀取內容並分塊...")
    try:
        with open(DB_DIR / "O.json", "r", encoding="utf-8") as f:
            processed_data = json.load(f)
        logger.info(f"從 db/O.json 載入 {len(processed_data)} 條記錄。")
    except FileNotFoundError:
        logger.error("❌ db/O.json 未找到，請確保 process_files.py 成功運行。")
        return "❌ 原始處理文件 db/O.json 未找到。請檢查 data 資料夾是否有文件且 process_files.py 能正常運行。"
    except json.JSONDecodeError as e:
        logger.error(f"❌ 解析 db/O.json 失敗：{e}", exc_info=True)
        return f"❌ 解析 db/O.json 失敗：{e}"
    
    separators = ['，', ',','。', '？', '！', '.', '?', '!', '\n\n']
    splitter = RecursiveCharacterTextSplitter(separators=separators,chunk_size=256, chunk_overlap=128)
    
    for item in processed_data:
        project_name = item.get("project", "general")
        source_name = item.get("filename", "unknown_file")
        content = item.get("content", "")

        # 判斷內容是否為實際文本，並進行分塊
        # process_files.py 會為圖片或不支援的檔案寫入描述性文字，我們也將其視為一個chunk
        if content and not content.startswith(("讀取檔案時發生錯誤", "不支援此類型的檔案讀取", "壓縮檔案")):
            # 如果內容是正常文本且不是錯誤/不支持訊息
            chunks = splitter.split_text(content)
            for chunk in chunks:
                all_chunks_with_metadata.append({"project": project_name, "source": source_name, "content": chunk})
            logger.debug(f"檔案 '{source_name}' 分塊為 {len(chunks)} 個片段。")
        elif content:
            # 如果是圖片或不支持的檔案類型，但有描述性內容，則將其作為一個完整的 chunk 加入
            all_chunks_with_metadata.append({"project": project_name, "source": source_name, "content": content})
            logger.debug(f"檔案 '{source_name}' 作為單一描述性片段處理。")
        else:
            logger.warning(f"檔案 '{source_name}' (專案 '{project_name}') 內容為空，已跳過。")


    if not all_chunks_with_metadata:
        logger.warning("在 db/O.json 中找不到可處理的內容，或內容為空。")
        return "⚠️ 在 db/O.json 中找不到可處理的內容，或內容為空。"
    
    progress(0.6, desc="步驟 3/5: 轉換文本為向量...")
    contents_to_embed = [chunk['content'] for chunk in all_chunks_with_metadata]
    embeddings = RAG_COMPONENTS["model"].encode(contents_to_embed, convert_to_numpy=True, show_progress_bar=True)
    
    progress(0.8, desc="步驟 4/5: 建立並訓練 FAISS 索引...")
    num_chunks, embedding_dim = embeddings.shape
    if num_chunks < 40:
        index = faiss.IndexFlatL2(embedding_dim)
        logger.info(f"選擇 IndexFlatL2 索引，數據量：{num_chunks}")
    elif num_chunks < 100000:
        nlist = int(4 * np.sqrt(num_chunks))
        quantizer = faiss.IndexFlatL2(embedding_dim)
        index = faiss.IndexIVFFlat(quantizer, embedding_dim, nlist, faiss.METRIC_L2)
        logger.info(f"開始訓練 IndexIVFFlat，數據量：{num_chunks}")
        index.train(embeddings)
        logger.info(f"IndexIVFFlat 訓練完成。")
        logger.info(f"選擇 IndexIVFFlat 索引，數據量：{num_chunks}, nlist={nlist}")
    else:
        nlist = int(4 * np.sqrt(num_chunks))
        m = 64
        nbits = 8
        quantizer = faiss.IndexFlatL2(embedding_dim)
        index = faiss.IndexIVFPQ(quantizer, embedding_dim, nlist, m, nbits)
        logger.info(f"開始訓練 IndexIVFPQ，數據量：{num_chunks}, nlist={nlist}, m={m}, nbits={nbits}")
        index.train(embeddings)
        logger.info(f"IndexIVFPQ 訓練完成。")
        logger.info(f"選擇 IndexIVFPQ 索引，數據量：{num_chunks}, nlist={nlist}, m={m}, nbits={nbits}")

    index.add(embeddings); progress(0.9, desc="步驟 5/5: 儲存索引與元資料...")
    faiss.write_index(index, str(DB_DIR / "faiss.index"))
    with open(DB_DIR / "meta.json", "w", encoding="utf-8") as f: json.dump(all_chunks_with_metadata, f, ensure_ascii=False, indent=4)
    
    RAG_COMPONENTS["index"] = index; RAG_COMPONENTS["metadata"] = all_chunks_with_metadata
    logger.info(f"✅ 知識庫建立完成！共索引了 {num_chunks} 個知識片段。")
    return f"✅ 知識庫建立完成！共索引了 {num_chunks} 個知識片段。"

def rerank_results(query, docs):
    pairs = [[query, d["content"]] for d in docs]
    scores = reranker.predict(pairs)
    scored_docs = sorted(zip(docs, scores), key=lambda x: x[1], reverse=True)
    return [doc for doc, score in scored_docs]

def get_chatbot_response(message_list):
    logger.info("開始處理聊天機器人請求。")
    if not all(RAG_COMPONENTS.values()):
        message_list.append({"role": "assistant", "content": "錯誤：知識庫尚未載入。請先點擊『建立/更新索引』"})
        logger.warning("知識庫尚未載入，無法回應聊天機器人請求。")
        return message_list
    
    user_message = message_list[-1]['content']
    conversation_history = message_list[:-1] # 取得除了最後一個問題之外的所有對話
    history_str = ""
    for msg in conversation_history:
        history_str += f"{msg['role']}: {msg['content']}\n"
    
    logger.info(f"使用者問題：'{user_message}'")
    logger.info(f"歷史對話：\n{history_str}")
    
    index, metadata, model = RAG_COMPONENTS["index"], RAG_COMPONENTS["metadata"], RAG_COMPONENTS["model"]
    if hasattr(index, 'nprobe'): index.nprobe = 16
    
    try:
        query_embedding = model.encode([user_message], convert_to_numpy=True)
        distances, indices = index.search(query_embedding, 15)
    except Exception as e:
        message_list.append({"role": "assistant", "content": f"檢索時發生錯誤：{e}"})
        logger.error(f"檢索時發生錯誤：{e}", exc_info=True)
        return message_list
    
    docs = [metadata[i] for i in indices[0] if i != -1]
    docs = rerank_results(user_message, docs)
    top_docs = docs[:5]

    if not docs:
        message_list.append({"role": "assistant", "content": "抱歉，在我的知識庫中找不到與您問題相關的資訊。"})
        logger.warning("找不到相關文件，無法回答。")
        return message_list
    
    context = "\n\n".join([doc["content"] for doc in top_docs])
    prompt = ChatPromptTemplate.from_template(
    """你是一個專業的服務人員。僅根據檢索到的資料中分析所需的資料並且盡可能地回答使用者問題。如果資料不足，請回答「資料不足，無法判斷」。

    --- 歷史對話 ---
    {history}

    --- 上下文 ---
    {context}

    --- 問題 ---
    {question}

    回答：""")
    
    llm = OllamaLLM(model=OLLAMA_MODEL)
    chain = prompt | llm
    # 使用 .stream() 來獲取一個生成器
    stream_generator = chain.stream({
        "history": history_str,
        "context": context,
        "question": user_message
    })

    # 在 Gradio 聊天視窗中新增一個空的回答，準備接收串流
    message_list.append({"role": "assistant", "content": ""})

    # 逐一處理 stream_generator 傳回的片段
    for chunk in stream_generator:
        # 將新片段加到最後一則訊息中
        message_list[-1]["content"] += chunk
        # 使用 yield 回傳當前完整的聊天歷史，Gradio 會即時更新介面
        yield message_list
    
    # 迴圈結束後，完整的回答已經在 chat_history 中
    # 現在將參考資料加到最後一則訊息中
    sources = sorted(list(set(doc["source"] for doc in docs)))
    message_list[-1]["content"] += f"\n\n---\n*參考資料：{', '.join(sources)}*"
    
    # 最後再 yield 一次，確保參考資料顯示出來
    yield message_list
    
    logger.info("成功生成聊天機器人回應。")

def generate_document_from_template(template_name, project_scope, progress=gr.Progress()):
    logger.info(f"開始根據範本生成文件：範本='{template_name}', 專案='{project_scope}'。")
    if not all(RAG_COMPONENTS.values()):
        logger.warning("知識庫尚未載入，無法生成文件。")
        return gr.File(visible=False), "錯誤：知識庫尚未載入。請先建立索引。"
    if not template_name or not project_scope:
        logger.warning("範本或專案範圍未指定。")
        return gr.File(visible=False), "錯誤：請先選擇範本和專案範圍。"
    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        logger.error(f"找不到範本檔案 '{template_path}'。")
        return gr.File(visible=False), f"錯誤：找不到範本檔案 {template_name}"

    progress(0.1, desc=f"正在分析範本 '{template_name}'...")
    doc = docx.Document(template_path)
    
    placeholders = set()
    pattern = re.compile(r'\{\{(.*?)\}\}')
    for p in doc.paragraphs: placeholders.update(pattern.findall(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    placeholders.update(pattern.findall(p.text))
    
    if not placeholders:
        logger.warning("在範本中沒有找到任何 {{...}} 格式的預留位置。")
        return gr.File(visible=False), "警告：在範本中沒有找到任何 {{...}} 格式的預留位置。"

    replacements = {}
    placeholders = list(placeholders)
    
    for i, ph in enumerate(placeholders):
        progress(0.1 + (0.8 * (i+1)/len(placeholders)), desc=f"正在研究 '{ph}'...")
        question = f"在這個專案中「{ph}」的資料是什麼"
        answer = query_single_fact(question, project_scope)
        replacements[ph] = answer
        logger.info(f"研究結果 -> {ph}: {answer}")

    progress(0.9, desc="數據收集完成，正在生成最終報告...")
    
    for p in doc.paragraphs:
        for key, value in replacements.items():
            p.text = p.text.replace(f"{{{{{key}}}}}", str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in replacements.items():
                        p.text = p.text.replace(f"{{{{{key}}}}}", str(value))
    
    output_filename = f"{project_scope}_{template_name}"
    output_path = OUTPUT_DIR / output_filename
    doc.save(output_path)
    logger.info(f"✅ 報告已生成並儲存於：'{output_path}'。共處理了 {len(placeholders)} 個欄位。")
    return gr.File(value=str(output_path), visible=True, label="下載通用報告"), f"✅ 報告已生成！共處理了 {len(placeholders)} 個欄位。"

def handle_upload(files):
    logger.info("開始處理上傳檔案。")
    if files is None:
        logger.warning("沒有選擇要上傳的檔案。")
        return "請選擇要上傳的檔案。"
    for file_or_path in files:
        source_path = Path(file_or_path.name)
        try:
            if source_path.suffix == '.zip':
                with zipfile.ZipFile(source_path, 'r') as zip_ref:
                    unzip_dir = SOURCE_DIR / source_path.stem; unzip_dir.mkdir(exist_ok=True, parents=True)
                    zip_ref.extractall(unzip_dir)
                logger.info(f"成功解壓縮 '{source_path}' 至 '{unzip_dir}'。")
            else:
                shutil.copy(source_path, SOURCE_DIR / source_path.name)
                logger.info(f"成功複製 '{source_path}' 至 '{SOURCE_DIR}'。")
        except Exception as e:
            logger.error(f"處理上傳檔案 '{source_path}' 失敗：{e}", exc_info=True)
            return f"❌ 上傳失敗：{e}"
    logger.info("檔案/資料夾上傳成功。")
    return "✅ 檔案/資料夾上傳成功！請點擊下方按鈕開始建立索引。"

def get_dynamic_choices():
    templates = [f.name for f in TEMPLATE_DIR.glob("*.docx")]
    projects = []
    if RAG_COMPONENTS["metadata"]:
        projects = sorted(list(set(item.get("project", "general") for item in RAG_COMPONENTS["metadata"])))
    else:
        projects = [d.name for d in SOURCE_DIR.iterdir() if d.is_dir()]
        if not projects: projects.append("general")
    logger.info(f"動態載入選單：範本={templates}, 專案={projects}")
    return gr.Dropdown(choices=templates or ["無範本"], interactive=True), gr.Dropdown(choices=projects or ["無專案"], interactive=True)


# --- 應用程式啟動與 Gradio 介面 ---
logger.info("--- 應用程式啟動：正在載入模型並嘗試載入現有知識庫... ---")
try:
    RAG_COMPONENTS["model"] = SentenceTransformer(MODEL_NAME, trust_remote_code=True)
    logger.info("SentenceTransformer 模型載入成功。")
    if (DB_DIR / "faiss.index").exists() and (DB_DIR / "meta.json").exists():
        RAG_COMPONENTS["index"] = faiss.read_index(str(DB_DIR / "faiss.index"))
        with open(DB_DIR / "meta.json", "r", encoding="utf-8") as f: RAG_COMPONENTS["metadata"] = json.load(f)
        if RAG_COMPONENTS["index"].d != RAG_COMPONENTS["model"].get_sentence_embedding_dimension():
            raise ValueError("載入的 FAISS 索引與模型維度不匹配。請重新建立索引。")
        logger.info("✅ 現有知識庫載入成功！")
    else:
        logger.info("ℹ️ 未找到現有知識庫。請上傳檔案後，點擊『建立/更新索引』。")
except Exception as e:
    logger.error(f"❌ 應用程式啟動失敗或知識庫載入失敗：{e}。請檢查設定並重新執行。", exc_info=True)
    RAG_COMPONENTS = {"model": None, "index": None, "metadata": None}

with gr.Blocks(title="智慧 RAG 助理", theme=gr.themes.Soft()) as demo:
    gr.Markdown("# 智慧 RAG 助理")
    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("## 控制台")
            with gr.Accordion("1. 更新知識庫", open=False):
                file_uploader = gr.Files(label="上傳檔案或包含資料夾的 .zip 檔", file_types=['.pdf', '.docx', '.xlsx', '.csv', '.txt', '.zip'])
                upload_status = gr.Textbox(interactive=False, show_label=False, placeholder="上傳狀態...")
                ingest_button = gr.Button("開始建立/更新索引", variant="primary")
                ingest_status = gr.Textbox(label="索引狀態", interactive=False, value="等待操作...", lines=1)
            with gr.Accordion("2. 根據模板生成文件", open=True):
                with gr.Row():
                    template_dropdown = gr.Dropdown(label="選擇文件範本", interactive=False)
                    project_dropdown = gr.Dropdown(label="選擇專案範圍", interactive=False)
                generate_button = gr.Button("生成文件", variant="primary")
                with gr.Row():
                    generate_status = gr.Textbox(label="生成狀態", interactive=False, lines=1, scale=2)
                    download_link = gr.File(label="下載", interactive=False, visible=False, scale=1)
        with gr.Column(scale=2):
            chatbot = gr.Chatbot(label="對話顯示區", height=500, type='messages')
            with gr.Row():
                msg_input = gr.Textbox(show_label=False, placeholder="在這裡輸入您的問題，然後按 Enter...", scale=4)
                submit_button = gr.Button("送出", variant="primary", scale=1)
    file_uploader.upload(fn=handle_upload, inputs=file_uploader, outputs=upload_status)
    ingest_button.click(fn=build_knowledge_base, outputs=ingest_status).then(fn=get_dynamic_choices, outputs=[template_dropdown, project_dropdown])
    generate_button.click(fn=generate_document_from_template, inputs=[template_dropdown, project_dropdown], outputs=[download_link, generate_status])
    
    def add_user_message_to_history(user_message, history):
        history.append({"role": "user", "content": user_message})
        return "", history
    msg_input.submit(add_user_message_to_history, [msg_input, chatbot], [msg_input, chatbot], queue=False).then(get_chatbot_response, chatbot, chatbot)
    submit_button.click(add_user_message_to_history, [msg_input, chatbot], [msg_input, chatbot], queue=False).then(get_chatbot_response, chatbot, chatbot)
    
    demo.load(fn=get_dynamic_choices, outputs=[template_dropdown, project_dropdown])

if __name__ == "__main__":
    demo.launch()